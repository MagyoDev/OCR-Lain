from io import BytesIO
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from PIL import Image, UnidentifiedImageError

from ocr_lain.config import OCRConfig
from ocr_lain.extractors.base import BaseExtractor
from ocr_lain.models import ExtractedBlock, ExtractResult
from ocr_lain.ocr_engine import OCREngine


class DOCXExtractor(BaseExtractor):

    SUPPORTED_EMBEDDED_IMAGE_EXTENSIONS = {
        ".png",
        ".jpg",
        ".jpeg",
        ".webp",
        ".bmp",
        ".tiff",
        ".tif",
    }

    def __init__(self, config: OCRConfig):
        self.config = config
        self.ocr_engine = OCREngine(config)

    def extract(self, file_path: Path) -> ExtractResult:

        blocks: list[ExtractedBlock] = []
        errors: list[str] = []

        try:
            document = Document(file_path)

            paragraph_text = self._extract_paragraph_text(document)

            if paragraph_text:
                blocks.append(
                    ExtractedBlock(
                        source_file=str(file_path),
                        source_type="docx",
                        location="document:paragraphs",
                        method="native_text",
                        text=paragraph_text,
                        metadata={},
                    )
                )

            table_texts = self._extract_table_texts(document)

            for table_index, table_text in enumerate(table_texts, start=1):
                if table_text:
                    blocks.append(
                        ExtractedBlock(
                            source_file=str(file_path),
                            source_type="docx",
                            location=f"table:{table_index}",
                            method="native_table_text",
                            text=table_text,
                            metadata={
                                "table": table_index,
                            },
                        )
                    )

            if self.config.ocr_embedded_images:
                image_blocks, image_errors = self._extract_embedded_images(file_path)

                blocks.extend(image_blocks)
                errors.extend(image_errors)

        except Exception as error:
            errors.append(f"Erro ao processar DOCX {file_path}: {error}")

        return ExtractResult(
            file_path=file_path,
            blocks=blocks,
            errors=errors,
        )

    def _extract_paragraph_text(self, document: Document) -> str:

        paragraphs: list[str] = []

        for paragraph in document.paragraphs:
            text = paragraph.text.strip()

            if text:
                paragraphs.append(text)

        return "\n\n".join(paragraphs)

    def _extract_table_texts(self, document: Document) -> list[str]:

        table_texts: list[str] = []

        for table in document.tables:
            rows: list[str] = []

            for row in table.rows:
                cell_values: list[str] = []

                for cell in row.cells:
                    cell_text = cell.text.strip().replace("\n", " ")

                    if cell_text:
                        cell_values.append(cell_text)

                if cell_values:
                    rows.append(" | ".join(cell_values))

            if rows:
                table_texts.append("\n".join(rows))

        return table_texts

    def _extract_embedded_images(
        self,
        file_path: Path,
    ) -> tuple[list[ExtractedBlock], list[str]]:


        blocks: list[ExtractedBlock] = []
        errors: list[str] = []

        try:
            with ZipFile(file_path) as docx_zip:
                media_files = [
                    name
                    for name in docx_zip.namelist()
                    if name.startswith("word/media/")
                    and Path(name).suffix.lower()
                    in self.SUPPORTED_EMBEDDED_IMAGE_EXTENSIONS
                ]

                for image_index, image_name in enumerate(media_files, start=1):
                    try:
                        image_data = docx_zip.read(image_name)
                        image = Image.open(BytesIO(image_data))

                        text = self.ocr_engine.image_to_text(image)

                        if text:
                            blocks.append(
                                ExtractedBlock(
                                    source_file=str(file_path),
                                    source_type="docx",
                                    location=f"embedded_image:{image_index}",
                                    method="ocr_embedded_image",
                                    text=text,
                                    metadata={
                                        "image": image_name,
                                        "image_index": image_index,
                                        "format": image.format,
                                        "width": image.width,
                                        "height": image.height,
                                    },
                                )
                            )

                    except UnidentifiedImageError:
                        errors.append(
                            f"Imagem interna não reconhecida em {file_path}: {image_name}"
                        )

                    except Exception as error:
                        errors.append(
                            f"Erro ao aplicar OCR na imagem {image_name} de {file_path}: {error}"
                        )

        except Exception as error:
            errors.append(f"Erro ao ler imagens internas do DOCX {file_path}: {error}")

        return blocks, errors